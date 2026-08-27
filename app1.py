import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import io
import zipfile
import requests


# ============================================================
# CONFIGURACIÓN DEL TAMAÑO DE LAS VARIABLES
# ============================================================

TAMANO_VARIABLE = 9


# ============================================================
# FUNCIONES AUXILIARES
# ============================================================

def limpiar_texto(texto):
    return (
        str(texto)
        .replace(" ", "_")
        .replace("/", "")
        .replace("\\", "")
    )


def aplicar_formato_variable(run, negrita=False, subrayado=False):
    """
    Formato que tendrán ÚNICAMENTE los datos reemplazados.
    """

    run.font.name = "Poppins"
    run.font.size = Pt(TAMANO_VARIABLE)
    run.bold = negrita
    run.underline = subrayado


def copiar_formato_run(run_origen, run_destino):
    """
    Copia el formato original de un run.
    """

    if run_origen._r.rPr is not None:
        run_destino._r.get_or_add_rPr().append(
            deepcopy(run_origen._r.rPr)
        )


# ============================================================
# REEMPLAZAR VARIABLE DENTRO DE UN SOLO RUN
# ============================================================

def reemplazar_variable_en_run(
    parrafo,
    run,
    variable,
    valor,
    negrita=False,
    subrayado=False
):

    if variable not in run.text:
        return False

    texto_original = run.text

    partes = texto_original.split(variable, 1)

    texto_antes = partes[0]
    texto_despues = partes[1]

    # Guardamos el formato original
    formato_original = None

    if run._r.rPr is not None:
        formato_original = deepcopy(run._r.rPr)

    # --------------------------------------------------------
    # El run original conservará el texto anterior
    # y su formato original
    # --------------------------------------------------------

    run.text = texto_antes

    # --------------------------------------------------------
    # Crear run con el valor reemplazado
    # --------------------------------------------------------

    run_valor = parrafo.add_run(str(valor))

    aplicar_formato_variable(
        run_valor,
        negrita=negrita,
        subrayado=subrayado
    )

    # Mover el run inmediatamente después del run original
    run._r.addnext(run_valor._r)

    # --------------------------------------------------------
    # Texto que estaba después de la variable
    # --------------------------------------------------------

    if texto_despues:

        run_despues = parrafo.add_run(texto_despues)

        if formato_original is not None:
            run_despues._r.insert(
                0,
                deepcopy(formato_original)
            )

        run_valor._r.addnext(run_despues._r)

    return True


# ============================================================
# REEMPLAZAR VARIABLE DIVIDIDA ENTRE VARIOS RUNS
# ============================================================

def reemplazar_variable_dividida(
    parrafo,
    variable,
    valor,
    negrita=False,
    subrayado=False
):

    runs = parrafo.runs

    if not runs:
        return False

    texto_completo = "".join(
        run.text for run in runs
    )

    posicion = texto_completo.find(variable)

    if posicion == -1:
        return False

    posicion_fin = posicion + len(variable)

    # --------------------------------------------------------
    # Determinar qué runs contienen la variable
    # --------------------------------------------------------

    acumulado = 0
    indices_afectados = []

    for i, run in enumerate(runs):

        inicio = acumulado
        fin = acumulado + len(run.text)

        if fin > posicion and inicio < posicion_fin:
            indices_afectados.append(i)

        acumulado = fin

    if not indices_afectados:
        return False

    primer_indice = indices_afectados[0]
    ultimo_indice = indices_afectados[-1]

    primer_run = runs[primer_indice]
    ultimo_run = runs[ultimo_indice]

    inicio_primer_run = sum(
        len(runs[i].text)
        for i in range(primer_indice)
    )

    inicio_ultimo_run = sum(
        len(runs[i].text)
        for i in range(ultimo_indice)
    )

    # --------------------------------------------------------
    # Texto antes y después de la variable
    # --------------------------------------------------------

    posicion_relativa_inicio = (
        posicion - inicio_primer_run
    )

    posicion_relativa_fin = (
        posicion_fin - inicio_ultimo_run
    )

    texto_antes = (
        primer_run.text[
            :posicion_relativa_inicio
        ]
    )

    texto_despues = (
        ultimo_run.text[
            posicion_relativa_fin:
        ]
    )

    # --------------------------------------------------------
    # Guardar formatos originales
    # --------------------------------------------------------

    formato_antes = None
    formato_despues = None

    if primer_run._r.rPr is not None:
        formato_antes = deepcopy(
            primer_run._r.rPr
        )

    if ultimo_run._r.rPr is not None:
        formato_despues = deepcopy(
            ultimo_run._r.rPr
        )

    # --------------------------------------------------------
    # Guardar posición XML del primer run
    # --------------------------------------------------------

    primer_elemento = primer_run._r
    padre = primer_elemento.getparent()

    # --------------------------------------------------------
    # Eliminar todos los runs que contienen la variable
    # --------------------------------------------------------

    elementos_eliminar = [
        runs[i]._r
        for i in indices_afectados
    ]

    for elemento in elementos_eliminar:
        padre.remove(elemento)

    # --------------------------------------------------------
    # Crear run del texto anterior
    # --------------------------------------------------------

    ultimo_insertado = None

    if texto_antes:

        run_antes = parrafo.add_run(texto_antes)

        if formato_antes is not None:
            run_antes._r.insert(
                0,
                deepcopy(formato_antes)
            )

        padre.insert(
            padre.index(primer_elemento)
            if primer_elemento.getparent() is padre
            else len(padre),
            run_antes._r
        )

        ultimo_insertado = run_antes._r

    # --------------------------------------------------------
    # Crear run del valor reemplazado
    # --------------------------------------------------------

    run_valor = parrafo.add_run(str(valor))

    aplicar_formato_variable(
        run_valor,
        negrita=negrita,
        subrayado=subrayado
    )

    # --------------------------------------------------------
    # Crear run del texto posterior
    # --------------------------------------------------------

    if texto_despues:

        run_despues = parrafo.add_run(
            texto_despues
        )

        if formato_despues is not None:
            run_despues._r.insert(
                0,
                deepcopy(formato_despues)
            )

    return True


# ============================================================
# REEMPLAZAR VARIABLES EN UN PÁRRAFO
# ============================================================

def reemplazar_variables_en_parrafo(
    parrafo,
    reemplazos
):

    for variable, datos in reemplazos.items():

        valor = datos["valor"]
        negrita = datos.get("negrita", False)
        subrayado = datos.get("subrayado", False)

        # ----------------------------------------------------
        # Intentar primero si la variable está en un solo run
        # ----------------------------------------------------

        encontrado = False

        for run in list(parrafo.runs):

            if variable in run.text:

                reemplazar_variable_en_run(
                    parrafo,
                    run,
                    variable,
                    valor,
                    negrita,
                    subrayado
                )

                encontrado = True
                break

        # ----------------------------------------------------
        # Si está dividida en varios runs
        # ----------------------------------------------------

        if not encontrado:

            reemplazar_variable_dividida(
                parrafo,
                variable,
                valor,
                negrita,
                subrayado
            )


# ============================================================
# PROCESAR TABLAS
# ============================================================

def procesar_tabla(tabla, reemplazos):

    for fila in tabla.rows:

        for celda in fila.cells:

            # Párrafos de la celda
            for parrafo in celda.paragraphs:

                reemplazar_variables_en_parrafo(
                    parrafo,
                    reemplazos
                )

            # Tablas dentro de tablas
            for tabla_interna in celda.tables:

                procesar_tabla(
                    tabla_interna,
                    reemplazos
                )


# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(
    page_title="Generador de Oficios",
    page_icon="📄",
    layout="wide"
)


# ============================================================
# ESTILO
# ============================================================

st.markdown(
    """
    <style>

    .stApp {
        background-color: #f6f8fb;
        color: #1f3b57;
    }

    h1, h2, h3 {
        color: #2a4d69;
    }

    .stButton>button {
        background: linear-gradient(
            135deg,
            #4c8bf5,
            #6fa8ff
        );

        color: white;
        border-radius: 8px;
        padding: 0.6em 1.2em;
        border: none;
    }

    .stButton>button:hover {
        background: linear-gradient(
            135deg,
            #3a73d9,
            #558eff
        );
    }

    </style>
    """,
    unsafe_allow_html=True
)


# ============================================================
# TÍTULO
# ============================================================

st.title("📄 Generador de Oficios")

st.info(
    """
    Plataforma para la generación automatizada de oficios
    institucionales utilizando una plantilla Word y datos
    provenientes de un directorio de empresas, permitiendo
    producir documentos oficiales de manera rápida, uniforme
    y eficiente.
    """
)


# ============================================================
# GOOGLE SHEETS
# ============================================================

SHEET_URL = (
    "https://script.google.com/macros/s/"
    "AKfycbwgeHn7HqMoFH4VkqyMRGZ8v-B7YAFA8_PgH1XYnIzHfwmSEIG"
    "aweuPFjjdbMFTjC_0rg/exec"
)


try:

    response = requests.get(
        SHEET_URL,
        timeout=30
    )

    if response.status_code != 200:

        st.error(
            "No se pudo conectar con Google Sheets"
        )

        st.stop()

    data = response.json()

    df = pd.DataFrame(data)

    df.columns = df.columns.str.strip()


except Exception as e:

    st.error(
        f"Error al conectar con Google Sheets: {e}"
    )

    st.stop()


# ============================================================
# EMPRESAS
# ============================================================

st.subheader("🏢 Seleccione empresas")


razones = (
    df["Nombre de la Empresa"]
    .dropna()
    .unique()
    .tolist()
)


razones_seleccionadas = st.multiselect(
    "Empresas disponibles:",
    razones
)


st.divider()


# ============================================================
# PLANTILLA WORD
# ============================================================

st.subheader("📄 Subir plantilla Word")


st.info(
    """
    Para que el sistema reemplace los datos automáticamente,
    copie y pegue exactamente las siguientes variables dentro
    del documento Word:
    """
)


st.code(
    """
Señor
[Nombre del Destinatario]
[Cargo]
[Entidad]
[Dirección]
[Distrito]

Asunto: [Asunto]
""",
    language="text"
)


word_file = st.file_uploader(
    "Plantilla Word (.docx)",
    type=["docx"]
)


st.divider()


# ============================================================
# DATOS ADICIONALES
# ============================================================

st.subheader("✏️ Información adicional")


asunto = st.text_input(
    "Asunto del oficio",
    placeholder="Ejemplo: Solicitud de información técnica"
)


procedimiento = st.text_input(
    "Procedimiento (solo para el nombre del archivo)",
    placeholder="Ejemplo: ERACMF"
)


st.divider()


# ============================================================
# ARCHIVOS ADJUNTOS
# ============================================================

st.subheader(
    "📎 Adjuntar archivos por tipo de actividad"
)


tipos_archivo = [
    "pdf",
    "docx",
    "doc",
    "xlsx",
    "xls",
    "txt",
    "zip",
    "rar",
    "csv"
]


col1, col2 = st.columns(2)


with col1:

    archivos_transmision = st.file_uploader(
        "Archivos para Transmisión",
        type=tipos_archivo,
        accept_multiple_files=True
    )

    archivos_distribucion = st.file_uploader(
        "Archivos para Distribución",
        type=tipos_archivo,
        accept_multiple_files=True
    )


with col2:

    archivos_generacion = st.file_uploader(
        "Archivos para Generación",
        type=tipos_archivo,
        accept_multiple_files=True
    )

    archivos_cliente_libre = st.file_uploader(
        "Archivos para Cliente Libre",
        type=tipos_archivo,
        accept_multiple_files=True
    )


archivos_por_actividad = {

    "Transmisión": archivos_transmision,

    "Generación": archivos_generacion,

    "Distribución": archivos_distribucion,

    "Cliente Libre": archivos_cliente_libre

}


st.divider()


# ============================================================
# GENERAR OFICIOS
# ============================================================

st.subheader("⚙️ Generar Oficios")


if st.button("🚀 Generar y descargar ZIP"):

    if not word_file or not razones_seleccionadas:

        st.warning(
            "Debe subir la plantilla Word y seleccionar empresas."
        )

        st.stop()


    zip_buffer = io.BytesIO()


    with zipfile.ZipFile(
        zip_buffer,
        "w",
        zipfile.ZIP_DEFLATED
    ) as zip_file:


        for razon in razones_seleccionadas:

            fila = df[
                df["Nombre de la Empresa"] == razon
            ]


            if fila.empty:
                continue


            # ------------------------------------------------
            # OBTENER DATOS
            # ------------------------------------------------

            nombre_destinatario = (
                fila["GERENTE GENERAL"].values[0]
            )

            cargo = (
                fila["CARGO DEL REPRESENTANTE"].values[0]
            )

            entidad = (
                fila["Nombre de la Empresa"].values[0]
            )

            direccion = (
                fila["DIRECCIÓN"].values[0]
            )

            distrito = (
                fila["Distrito"].values[0]
            )

            actividad = (
                fila["ACTIVIDAD"].values[0]
            )

            codigo = (
                fila["CODIGO"].values[0]
            )


            # ------------------------------------------------
            # ABRIR PLANTILLA
            # ------------------------------------------------

            documento = Document(word_file)


            # ------------------------------------------------
            # VARIABLES A REEMPLAZAR
            # ------------------------------------------------

            reemplazos = {

                "[Nombre del Destinatario]": {
                    "valor": nombre_destinatario,
                    "negrita": True,
                    "subrayado": False
                },

                "[Cargo]": {
                    "valor": cargo,
                    "negrita": False,
                    "subrayado": False
                },

                "[Entidad]": {
                    "valor": entidad,
                    "negrita": True,
                    "subrayado": False
                },

                "[Dirección]": {
                    "valor": direccion,
                    "negrita": False,
                    "subrayado": False
                },

                "[Distrito]": {
                    "valor": distrito,
                    "negrita": False,
                    "subrayado": True
                },

                "[Asunto]": {
                    "valor": asunto,
                    "negrita": False,
                    "subrayado": False
                }

            }


            # ------------------------------------------------
            # PROCESAR PÁRRAFOS NORMALES
            # ------------------------------------------------

            for parrafo in documento.paragraphs:

                reemplazar_variables_en_parrafo(
                    parrafo,
                    reemplazos
                )


            # ------------------------------------------------
            # PROCESAR TABLAS
            # ------------------------------------------------

            for tabla in documento.tables:

                procesar_tabla(
                    tabla,
                    reemplazos
                )


            # ------------------------------------------------
            # GUARDAR DOCUMENTO
            # ------------------------------------------------

            doc_buffer = io.BytesIO()

            documento.save(
                doc_buffer
            )

            doc_buffer.seek(0)


            # ------------------------------------------------
            # NOMBRE DEL DOCUMENTO
            # ------------------------------------------------

            empresa_limpia = limpiar_texto(
                entidad
            )

            asunto_limpio = limpiar_texto(
                asunto
            )

            procedimiento_limpio = limpiar_texto(
                procedimiento
            )


            nombre_documento = (
                f"Oficio_"
                f"{procedimiento_limpio}_"
                f"{empresa_limpia}_"
                f"{asunto_limpio}.docx"
            )


            # ------------------------------------------------
            # RUTA DENTRO DEL ZIP
            # ------------------------------------------------

            ruta_doc = (
                f"{actividad}/"
                f"{codigo}/"
                f"{nombre_documento}"
            )


            zip_file.writestr(
                ruta_doc,
                doc_buffer.read()
            )


            # ------------------------------------------------
            # ARCHIVOS ADJUNTOS
            # ------------------------------------------------

            archivos = archivos_por_actividad.get(
                actividad,
                []
            )


            if archivos:

                for archivo in archivos:

                    zip_file.writestr(
                        f"{actividad}/"
                        f"{codigo}/"
                        f"{archivo.name}",
                        archivo.read()
                    )


    # ========================================================
    # DESCARGA DEL ZIP
    # ========================================================

    zip_buffer.seek(0)


    st.success(
        "✅ Oficios generados correctamente"
    )


    st.download_button(
        label="📦 Descargar ZIP",
        data=zip_buffer,
        file_name="oficios_generados.zip",
        mime="application/zip"
    )
